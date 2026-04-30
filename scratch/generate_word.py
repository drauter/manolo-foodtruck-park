from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('Informe de Defensa de Proyecto: Ecosistema Digital Manolo Foodtruck Park', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    doc.add_paragraph(
        'El proyecto Manolo Foodtruck Park no es solo un software de ventas, sino una solución integral '
        'de gestión operativa y financiera diseñada específicamente para la dinámica de alta velocidad '
        'de un parque de comida. A través de un ecosistema que conecta en tiempo real a vendedores, '
        'administradores y clientes finales, el sistema optimiza el flujo de trabajo, reduce errores '
        'humanos y proporciona una visibilidad financiera sin precedentes.'
    )

    # Core Pillars
    doc.add_heading('2. Pilares de la Solución', level=1)
    
    doc.add_heading('A. Omnicanalidad y Experiencia del Cliente', level=2)
    p1 = doc.add_paragraph()
    p1.add_run('Seller POS (Punto de Venta del Vendedor): ').bold = True
    p1.add_run('Interfaz optimizada para rapidez, permitiendo procesar pedidos en segundos con soporte para pagos complejos.')
    
    p2 = doc.add_paragraph()
    p2.add_run('Client Menu (Auto-servicio): ').bold = True
    p2.add_run('Menú digital accesible vía QR, empoderando al cliente y reduciendo la carga sobre el personal.')
    
    p3 = doc.add_paragraph()
    p3.add_run('Order Tracking & Public Display: ').bold = True
    p3.add_run('Sistema de monitoreo en tiempo real con anuncios de voz automatizados que informan al cliente exactamente cuándo su pedido está listo.')

    doc.add_heading('B. Gestión Administrativa Avanzada', level=2)
    p4 = doc.add_paragraph()
    p4.add_run('Analítica Predictiva: ').bold = True
    p4.add_run('Integración de Recharts para visualizar tendencias de ventas, picos de demanda y proyecciones de stock.')
    
    p5 = doc.add_paragraph()
    p5.add_run('Control de Deudas y Créditos: ').bold = True
    p5.add_run('Módulo especializado para gestionar cuentas por cobrar y historial de pagos.')

    doc.add_heading('C. Excelencia Técnica e Integración de Hardware', level=2)
    p6 = doc.add_paragraph()
    p6.add_run('Impresión Térmica Robusta: ').bold = True
    p6.add_run('Integración con QZ Tray para impresión de tickets en papel de 80mm.')
    
    p7 = doc.add_paragraph()
    p7.add_run('Arquitectura Real-Time: ').bold = True
    p7.add_run('Basado en Supabase, garantizando sincronización instantánea.')

    # Business Model
    doc.add_heading('3. Estrategia de Negocio: Modelo SaaS (Software as a Service)', level=1)
    doc.add_paragraph(
        'El sistema opera bajo un modelo de suscripción (SaaS) para garantizar sostenibilidad y '
        'soporte continuo. Argumentos clave:'
    )
    
    bullets = [
        ('Valor Continuo:', 'Garantiza actualizaciones, parches de seguridad y mejoras constantes sin costos extras.'),
        ('Inversión Inteligente:', 'Costo inicial bajo que facilita la entrada a nuevos locatarios sin inversión pesada en infraestructura.'),
        ('Mantenimiento:', 'Nos encargamos de servidores, copias de seguridad y soporte 24/7.'),
        ('Seguridad de Datos:', 'Información blindada en la nube, protegida contra fallos físicos en el local.')
    ]
    
    for bold_text, normal_text in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(' ' + normal_text)

    # Conclusion
    doc.add_heading('4. Conclusión', level=1)
    doc.add_paragraph(
        'Manolo Foodtruck Park transforma un negocio tradicional en una operación inteligente. '
        'Es una infraestructura crítica que permite escalar, controlar finanzas y ofrecer '
        'una experiencia de clase mundial.'
    )

    doc.save('Informe_Defensa_Manolo_Foodtruck.docx')
    print("Archivo generado exitosamente.")

if __name__ == "__main__":
    create_report()
